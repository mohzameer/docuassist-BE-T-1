from typing import List, Optional, Dict, Any, Tuple
from datetime import datetime
import logging
import io
import os
import tempfile
import PyPDF2
from docx import Document
from PIL import Image, UnidentifiedImageError
import pytesseract
from app.models.document import DocumentResponse
from app.services.google_drive_service import GoogleDriveService, GOOGLE_MIME_TYPES
from app.services.search_service import SearchService
from app.models.batch import BatchProcessingStatus, BatchProcessingStats, BatchProcessingError
import uuid
import asyncio
import re
from llama_index.core import Document, Settings
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import TextNode
from app.core.config import get_settings

logger = logging.getLogger(__name__)

SUPPORTED_MIME_TYPES = {
    'text/plain',
    'text/csv',
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'application/msword',
    'image/jpeg',
    'image/png',
    'image/gif',
    'image/tiff',
    # Add Google Workspace MIME types
    'application/vnd.google-apps.document',
    'application/vnd.google-apps.spreadsheet',
    'application/vnd.google-apps.presentation',
    'application/vnd.google-apps.drawing',
}

settings = get_settings()

# Configure LlamaIndex settings
Settings.chunk_size = settings.CHUNK_SIZE
Settings.chunk_overlap = settings.CHUNK_OVERLAP

class DocumentService:
    def __init__(
        self,
        drive_service: GoogleDriveService,
        search_service: SearchService
    ):
        self.drive_service = drive_service
        self.search_service = search_service
        self._batch_status = {}
        
        # Initialize text splitter with LlamaIndex
        self.text_splitter = SentenceSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            paragraph_separator="\n\n",
            secondary_chunking_regex="[^,.;]+[,.;]?",
        )

    async def _process_pdf(self, file_content: bytes, filename: str) -> str:
        """Extract text from PDF file"""
        logger.info(f"Processing PDF file: {filename}")
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page_num, page in enumerate(pdf_reader.pages, 1):
            logger.debug(f"Processing page {page_num} of {filename}")
            text += page.extract_text() + "\n"
        return text

    async def _process_docx(self, file_content: bytes, filename: str) -> str:
        """Extract text from DOCX file"""
        logger.info(f"Processing DOCX file: {filename}")
        doc = Document(io.BytesIO(file_content))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    async def _process_image(self, file_content: bytes, filename: str) -> str:
        """Extract text from image using OCR"""
        logger.info(f"Processing image file: {filename}")
        try:
            # Create a temporary file to save the image
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                try:
                    # Open image from bytes
                    with Image.open(io.BytesIO(file_content)) as image:
                        logger.debug(f"Image format: {image.format}, Mode: {image.mode}, Size: {image.size}")
                        
                        # Convert image to RGB mode if necessary
                        if image.mode not in ('L', 'RGB'):
                            logger.debug(f"Converting image mode from {image.mode} to RGB")
                            image = image.convert('RGB')
                        
                        # Save the image as PNG
                        image.save(temp_file.name, 'PNG')
                        logger.debug(f"Saved temporary image: {temp_file.name}")
                    
                    # Use pytesseract to extract text
                    text = pytesseract.image_to_string(temp_file.name)
                    extracted_text = text.strip()
                    logger.debug(f"Extracted {len(extracted_text)} characters from image")
                    return extracted_text
                    
                finally:
                    # Clean up the temporary file
                    try:
                        os.unlink(temp_file.name)
                        logger.debug(f"Cleaned up temporary file: {temp_file.name}")
                    except Exception as e:
                        logger.warning(f"Error cleaning up temporary file for {filename}: {e}")
            
        except UnidentifiedImageError:
            logger.error(f"Could not identify image format for file: {filename}")
            raise
        except Exception as e:
            logger.error(f"Error processing image {filename}: {str(e)}")
            raise

    def _extract_metadata(self, document: DocumentResponse, content: str) -> Dict[str, Any]:
        """Extract rich metadata from document"""
        try:
            # Basic metadata
            metadata = {
                "document_id": document.id,
                "title": document.title,
                "drive_id": document.drive_id,
                "mime_type": document.mime_type,
                "processed_timestamp": str(datetime.utcnow()),
                "is_google_workspace": document.mime_type in GOOGLE_MIME_TYPES,
            }

            # Extract sections/headers
            sections = self._extract_sections(content)
            if sections:
                metadata["sections"] = sections

            # Extract document structure
            structure = self._analyze_document_structure(content)
            if structure:
                metadata["document_structure"] = structure

            return metadata
        except Exception as e:
            logger.error(f"Error extracting metadata: {e}")
            return metadata

    def _extract_sections(self, content: str) -> List[Dict[str, Any]]:
        """Extract document sections based on headers"""
        sections = []
        lines = content.split('\n')
        current_section = None
        
        for line in lines:
            # Identify headers (simple heuristic)
            stripped = line.strip()
            if stripped and (
                all(c.isupper() for c in stripped) or  # ALL CAPS
                stripped.startswith('#') or            # Markdown style
                re.match(r'^[0-9]+\.\s', stripped)     # Numbered sections
            ):
                if current_section:
                    sections.append(current_section)
                current_section = {
                    "title": stripped,
                    "content": "",
                    "start_position": len('\n'.join(lines[:lines.index(line)]))
                }
            elif current_section:
                current_section["content"] += line + "\n"

        if current_section:
            sections.append(current_section)
        return sections

    def _analyze_document_structure(self, content: str) -> Dict[str, Any]:
        """Analyze document structure for better context"""
        return {
            "total_length": len(content),
            "paragraphs": len([p for p in content.split('\n\n') if p.strip()]),
            "sentences": len([s for s in content.split('.') if s.strip()]),
            "has_tables": bool(re.search(r'\|[-+]+\|', content)),  # Simple table detection
            "has_lists": bool(re.search(r'^\s*[-*•]\s', content, re.MULTILINE)),  # List detection
        }

    def _extract_surrounding_context(
        self,
        chunk: str,
        full_text: str,
        context_window: int = 200
    ) -> Dict[str, str]:
        """Extract context around a chunk"""
        try:
            chunk_start = full_text.find(chunk)
            if chunk_start == -1:
                return {}

            # Get surrounding text
            context_start = max(0, chunk_start - context_window)
            context_end = min(len(full_text), chunk_start + len(chunk) + context_window)

            # Find sentence boundaries
            prev_period = full_text.rfind('.', context_start, chunk_start)
            next_period = full_text.find('.', chunk_start + len(chunk), context_end)

            if prev_period != -1:
                context_start = prev_period + 1
            if next_period != -1:
                context_end = next_period + 1

            return {
                "previous_context": full_text[context_start:chunk_start].strip(),
                "next_context": full_text[chunk_start + len(chunk):context_end].strip(),
                "position": chunk_start,
                "relative_position": chunk_start / len(full_text)
            }
        except Exception as e:
            logger.error(f"Error extracting context: {e}")
            return {}

    def _clean_text(self, text: str) -> str:
        """Enhanced text cleaning"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Normalize quotes and dashes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace('–', '-').replace('—', '-')
        
        # Fix common OCR errors
        text = text.replace('|', 'I').replace('0', 'O')
        
        # Remove special characters but keep meaningful punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\-\(\)\[\]\{\}\'\"\?\/\&\%\$\@\!]', '', text)
        
        # Normalize newlines
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()

    async def _process_single_document(
        self,
        document: DocumentResponse
    ) -> Tuple[Optional[str], Dict[str, Any]]:
        """Process a single document with enhanced context"""
        logger.info(f"Processing document: {document.title} (ID: {document.id})")
        
        try:
            # Skip unsupported file types
            if document.mime_type not in SUPPORTED_MIME_TYPES:
                logger.warning(f"Unsupported MIME type: {document.mime_type}")
                return None, {}

            # Download file from Google Drive
            file_content, mime_type = await self.drive_service.download_file(
                document.drive_id,
                document.mime_type
            )

            # Extract text based on mime type
            if mime_type == 'application/pdf':
                content = await self._process_pdf(file_content, document.title)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                content = await self._process_docx(file_content, document.title)
            elif mime_type.startswith('image/'):
                content = await self._process_image(file_content, document.title)
            else:
                content = file_content.decode('utf-8', errors='ignore')

            if not content:
                logger.warning(f"No content extracted from document: {document.title}")
                return None, {}

            # Clean the text
            cleaned_content = self._clean_text(content)

            # Extract rich metadata
            enhanced_metadata = self._extract_metadata(document, cleaned_content)

            # Create nodes with context
            nodes = self._chunk_text(cleaned_content, enhanced_metadata)
            
            # Add context to each node
            processed_nodes = []
            for node in nodes:
                context = self._extract_surrounding_context(
                    node.text,
                    cleaned_content
                )
                node.metadata.update({
                    "context": context,
                    "section": self._find_section(node.text, enhanced_metadata.get("sections", [])),
                    "document_structure": enhanced_metadata["document_structure"]
                })
                processed_nodes.append(node)

            return processed_nodes, enhanced_metadata

        except Exception as e:
            logger.error(f"Error processing document {document.title}: {str(e)}")
            raise

    def _find_section(self, chunk_text: str, sections: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
        """Find which section a chunk belongs to"""
        if not sections:
            return None
            
        chunk_position = chunk_text.find(chunk_text)
        for section in sections:
            section_end = section["start_position"] + len(section["content"])
            if section["start_position"] <= chunk_position <= section_end:
                return {
                    "title": section["title"],
                    "relative_position": (chunk_position - section["start_position"]) / len(section["content"])
                }
        return None

    def _chunk_text(self, text: str, metadata: Dict[str, Any]) -> List[TextNode]:
        """Split text into chunks using LlamaIndex's text splitter"""
        try:
            # Clean the text
            text = self._clean_text(text)
            
            # Create a Document object with metadata
            doc = Document(
                text=text,
                metadata=metadata,
                excluded_metadata_keys=["chunk_id", "chunk_index"]
            )
            
            # Split into nodes using the correct method
            nodes = self.text_splitter.get_nodes_from_documents([doc])
            
            logger.info(f"Split text into {len(nodes)} chunks")
            return nodes
            
        except Exception as e:
            logger.error(f"Error chunking text: {e}")
            raise

    async def process_documents(
        self,
        documents: List[DocumentResponse]
    ) -> List[DocumentResponse]:
        """Process multiple documents and update search index"""
        logger.info(f"Starting batch processing of {len(documents)} documents")
        processed_docs = []
        failed_docs = []
        
        for document in documents:
            try:
                logger.info(f"Processing document: {document.title}")
                content, metadata = await self._process_single_document(document)
                if content:
                    # Create base metadata
                    base_metadata = {
                        **metadata,
                        "document_id": document.id,
                        "title": document.title,
                        "drive_id": document.drive_id
                    }
                    
                    # Split content into nodes
                    nodes = self._chunk_text(content, base_metadata)
                    
                    # Index each node
                    for i, node in enumerate(nodes):
                        # Add chunk-specific metadata
                        node.metadata.update({
                            "chunk_id": f"{document.id}_chunk_{i}",
                            "chunk_index": i,
                            "total_chunks": len(nodes)
                        })
                        
                        # Update search index
                        await self.search_service.update_index([{
                            "text": node.text,
                            "metadata": node.metadata
                        }])
                    
                    processed_docs.append(document)
                    logger.info(f"Successfully indexed document: {document.title} ({len(nodes)} chunks)")
                else:
                    failed_docs.append(document.title)
            except Exception as e:
                failed_docs.append(document.title)
                logger.error(f"Failed to process document {document.title}: {str(e)}")
                continue
        
        logger.info(f"Batch processing completed. Processed: {len(processed_docs)}, Failed: {len(failed_docs)}")
        if failed_docs:
            logger.info(f"Failed documents: {', '.join(failed_docs)}")
        
        return processed_docs

    async def start_batch_processing(
        self,
        documents: List[DocumentResponse]
    ) -> BatchProcessingStatus:
        """Start batch processing of documents"""
        batch_id = str(uuid.uuid4())
        total = len(documents)
        
        # Create initial status with stats
        self._batch_status[batch_id] = BatchProcessingStatus(
            batch_id=batch_id,
            status="processing",
            stats=BatchProcessingStats(
                total_documents=total,
                processed_documents=0,
                failed_documents=0
            )
        )
        
        # Start processing in background
        asyncio.create_task(self._process_batch(batch_id, documents))
        
        return self._batch_status[batch_id]

    async def _process_batch(
        self,
        batch_id: str,
        documents: List[DocumentResponse]
    ):
        """Process documents in batch"""
        try:
            for document in documents:
                try:
                    content, metadata = await self._process_single_document(document)
                    if content:
                        await self.search_service.update_index([{
                            "text": content,
                            "metadata": {
                                **metadata,
                                "document_id": document.id,
                                "title": document.title,
                                "drive_id": document.drive_id
                            }
                        }])
                        self._batch_status[batch_id].stats.processed_documents += 1
                    else:
                        self._batch_status[batch_id].stats.failed_documents += 1
                        self._batch_status[batch_id].errors.append(
                            BatchProcessingError(
                                document_id=document.id,
                                error_message="No content extracted"
                            )
                        )
                except Exception as e:
                    self._batch_status[batch_id].stats.failed_documents += 1
                    self._batch_status[batch_id].errors.append(
                        BatchProcessingError(
                            document_id=document.id,
                            error_message=str(e)
                        )
                    )
                    
            self._batch_status[batch_id].status = "completed"
            self._batch_status[batch_id].stats.end_time = datetime.utcnow()
            
        except Exception as e:
            logger.error(f"Batch processing error: {e}", exc_info=True)
            self._batch_status[batch_id].status = "failed"
            self._batch_status[batch_id].stats.end_time = datetime.utcnow()

    async def get_batch_status(self, batch_id: str) -> Optional[BatchProcessingStatus]:
        """Get status of a batch processing job"""
        return self._batch_status.get(batch_id)

    async def get_processing_progress(self, batch_id: str):
        """Stream batch processing progress"""
        while True:
            status = self._batch_status.get(batch_id)
            if not status:
                break
                
            yield status
            
            if status.status in ["completed", "failed"]:
                break
                
            await asyncio.sleep(1) 